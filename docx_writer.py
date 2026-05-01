"""
AI Job Agent - DOCX Resume Writer
Saves AI-tailored resumes as professional .docx files.
"""

import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Dict

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import OUTPUT_DIR

logger = logging.getLogger(__name__)


def _sanitize_filename(text: str) -> str:
    """Remove special characters for safe filenames."""
    return re.sub(r'[\\/*?:"<>|]', "", text).strip().replace(" ", "_")


def save_resume_as_docx(
    resume_text: str,
    job: Dict,
    output_dir: str = OUTPUT_DIR,
) -> str:
    """
    Save a tailored resume as a nicely formatted .docx file.

    Args:
        resume_text: The AI-rewritten resume plain text.
        job: Dict with title, company, date_posted, job_url.
        output_dir: Directory to save files in.

    Returns:
        Absolute path to the saved .docx file.
    """
    Path(output_dir).mkdir(parents=True, exist_ok=True)

    # Build filename
    company_safe = _sanitize_filename(job.get("company", "Company"))
    title_safe = _sanitize_filename(job.get("title", "Role"))
    date_str = datetime.today().strftime("%Y-%m-%d")
    filename = f"ATS_Resume_{company_safe}_{title_safe}_{date_str}.docx"
    filepath = os.path.join(output_dir, filename)

    doc = Document()

    # ── Page margins ──────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # ── Header banner ─────────────────────────────────────────
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run(f"ATS-Tailored Resume")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(
        f"{job.get('title', '')}  ·  {job.get('company', '')}  ·  Generated {date_str}"
    )
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # ── Divider ───────────────────────────────────────────────
    doc.add_paragraph("─" * 80)

    # ── Resume body ───────────────────────────────────────────
    lines = resume_text.splitlines()
    for line in lines:
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph()  # blank line
            continue

        # Detect section headers (ALL CAPS lines or lines ending in colon)
        is_header = (
            stripped.isupper()
            or (stripped.endswith(":") and len(stripped.split()) <= 4)
            or stripped in (
                "SUMMARY", "OBJECTIVE", "EXPERIENCE", "WORK EXPERIENCE",
                "EDUCATION", "SKILLS", "CERTIFICATIONS", "PROJECTS",
                "ACHIEVEMENTS", "AWARDS", "LANGUAGES", "PUBLICATIONS",
            )
        )

        if is_header:
            para = doc.add_paragraph()
            run = para.add_run(stripped)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)  # Blue
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after = Pt(2)
        else:
            para = doc.add_paragraph(stripped)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(1)
            for run in para.runs:
                run.font.size = Pt(10)

    # ── Footer with job link ──────────────────────────────────
    doc.add_paragraph("─" * 80)
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(f"Job URL: {job.get('job_url', 'N/A')}")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(filepath)
    logger.info(f"  ✓ Saved: {filepath}")
    return os.path.abspath(filepath)


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)

    sample_text = """
JOHN DOE
john@example.com | linkedin.com/in/johndoe | (555) 123-4567

SUMMARY
Results-driven Backend Engineer with 5+ years of experience building high-performance APIs.

EXPERIENCE
Senior Software Engineer — Acme Corp (2020–Present)
- Architected REST APIs serving 1M+ daily requests using Python and FastAPI
- Led migration to microservices reducing p99 latency by 40%
- Mentored 3 junior engineers across 2 teams

SKILLS
Python, FastAPI, Docker, Kubernetes, PostgreSQL, AWS, Redis

EDUCATION
B.S. Computer Science — State University, 2019
"""

    sample_job = {
        "title": "Senior Backend Engineer",
        "company": "TechStartup",
        "job_url": "https://linkedin.com/jobs/view/123",
        "date_posted": "2026-04-08",
    }

    path = save_resume_as_docx(sample_text, sample_job)
    print(f"Saved to: {path}")
